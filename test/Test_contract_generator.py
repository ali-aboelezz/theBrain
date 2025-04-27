import pytest
from app.core import ContractGenerator
from docx import Document


# Mock setup for testing
@pytest.fixture
def generator():
    return ContractGenerator()


@pytest.fixture
def sample_template(tmp_path):
    """Create a mock Word document for testing."""
    file_path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("This contract begins on [DATE] and involves [DISCLOSING_PARTY_NAME].")
    doc.save(file_path)
    return str(file_path)


def test_extract_placeholders(generator, sample_template):
    """Test the extraction of placeholders from a template."""
    placeholders = generator.extract_placeholders(sample_template)
    assert placeholders == ["DATE", "DISCLOSING_PARTY_NAME"]


def test_generate_questions(generator):
    """Test the generation of questions for placeholders."""
    placeholders = ["DATE", "DISCLOSING_PARTY_NAME"]
    questions = generator.generate_questions(placeholders)
    assert questions == {
        "DATE": "What is the effective date of the contract?",
        "DISCLOSING_PARTY_NAME": "Who is the disclosing party mentioned in the contract?"
    }


def test_fill_document(generator, sample_template, tmp_path):
    """Test filling placeholders in a document."""
    responses = {"DATE": "2025-04-12", "DISCLOSING_PARTY_NAME": "Acme Corporation"}
    output_path = tmp_path / "filled_template.docx"
    generator.fill_document(sample_template, str(output_path), responses)

    # Verify the document content
    doc = Document(str(output_path))
    text = [para.text for para in doc.paragraphs]
    assert "This contract begins on 2025-04-12 and involves Acme Corporation." in text