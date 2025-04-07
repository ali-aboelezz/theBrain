import argparse
import logging
import pandas as pd
from docx import Document
from transformers import pipeline
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Configure logging
logging.basicConfig(
    filename="contract_generator.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

class ContractGenerator:
    def __init__(self, template_path, output_path, contract_type):
        self.template_path = template_path
        self.output_path = output_path
        self.contract_type = contract_type
        self.placeholders = []
        self.responses = {}

        # Load Hugging Face model for question generation (using Flan-T5)
        try:
            logging.info("Loading Hugging Face Flan-T5 question generator model...")
            self.question_generator = pipeline("text2text-generation", model="google/flan-t5-base")
        except Exception as e:
            logging.error(f"Error loading model: {e}")
            raise

    def extract_placeholders(self):
        """ Extract placeholders from the template document (e.g., [FIELD_NAME]). """
        try:
            logging.info(f"Extracting placeholders from: {self.template_path}")
            document = Document(self.template_path)
            for para in document.paragraphs:
                if "[" in para.text and "]" in para.text:
                    field = para.text[para.text.find("[") + 1:para.text.find("]")]
                    self.placeholders.append(field)
            logging.info(f"Detected placeholders: {self.placeholders}")
        except Exception as e:
            logging.error(f"Error extracting placeholders: {e}")
            raise

    def generate_questions(self):
        """ Generate contextual and professional questions for each placeholder. """
        try:
            logging.info("Generating questions for placeholders...")
            questions = {}
            for placeholder in self.placeholders:
                context_question = self._generate_contextual_question(placeholder)

                # If no predefined question is found, dynamically generate one
                if context_question == "FALLBACK":
                    input_text = f"Generate a formal and professional question in English about the field '{placeholder}' in a contract."
                    generated_question = self.question_generator(input_text)[0]['generated_text']
                    questions[placeholder] = self._postprocess_question(generated_question)
                else:
                    questions[placeholder] = context_question
            logging.info("Generated questions successfully.")
            return questions
        except Exception as e:
            logging.error(f"Error generating questions: {e}")
            raise

    def _generate_contextual_question(self, placeholder):
        """ Return pre-defined professional questions, or fallback to model generation. """
        placeholder = placeholder.strip().upper()
        question_map = {
            'DATE': "What is the effective date of the contract?",
            'DISCLOSING_PARTY_NAME': "Who is the disclosing party mentioned in the contract?",
            'RECEIVING_PARTY_NAME': "Who is the receiving party of the confidential information?",
            'CONFIDENTIAL_INFO_DESCRIPTION': "What information is classified as confidential under this agreement?",
            'DURATION': "What is the duration of the agreement?"
        }

        return question_map.get(placeholder, "FALLBACK")

    def _postprocess_question(self, question):
        """ Post-processes generated questions to ensure professionalism. """
        question = question.strip()
        if question.startswith("What is the best way to describe"):
            question = question.replace("What is the best way to describe", "What is")
        question = question[0].capitalize() + question[1:]  # Capitalize first letter
        return question

    def collect_responses(self, questions):
        """ Collect user responses for placeholders dynamically. """
        try:
            logging.info("Collecting responses from the user...")
            for field, question in questions.items():
                while True:
                    user_input = input(f"{question}: ").strip()
                    if user_input:
                        self.responses[field] = user_input
                        break
                    else:
                        print("Input cannot be empty. Please try again.")
            logging.info(f"Collected responses: {self.responses}")
        except Exception as e:
            logging.error(f"Error collecting user responses: {e}")
            raise

    def fill_document(self, custom_output_path=None):
        """ Fill placeholders in the document with user responses. """
        try:
            logging.info(f"Filling placeholders in: {self.template_path}")
            document = Document(self.template_path)

            for para in document.paragraphs:
                for field, response in self.responses.items():
                    if f"[{field}]" in para.text:
                        para.text = para.text.replace(f"[{field}]", str(response))

            save_path = custom_output_path if custom_output_path else self.output_path
            document.save(save_path)
            logging.info(f"Contract saved as: {save_path}")
        except Exception as e:
            logging.error(f"Error filling document: {e}")
            raise

    def export_to_pdf(self, pdf_path):
        """ Export the user's answers to a PDF file. """
        try:
            logging.info(f"Exporting user responses to PDF: {pdf_path}")
            c = canvas.Canvas(pdf_path, pagesize=letter)

            c.setFont("Helvetica-Bold", 14)
            c.drawString(100, 750, "User Responses for Generated Contract")
            c.setFont("Helvetica", 12)

            y_pos = 730
            for field, response in self.responses.items():
                if y_pos < 50:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y_pos = 750
                c.drawString(100, y_pos, f"{field}: {response}")
                y_pos -= 20

            c.save()
            logging.info(f"PDF exported successfully to {pdf_path}")
        except Exception as e:
            logging.error(f"Error exporting to PDF: {e}")
            raise

    def handle_multiple_documents(self):
        """ Handle the creation of multiple documents using an Excel file. """
        try:
            logging.info("Handling multiple document generation...")

            self.extract_placeholders()

            excel_path = "contract_data.xlsx"
            logging.info(f"Creating Excel template at: {excel_path}")
            df = pd.DataFrame(columns=self.placeholders)
            df.to_excel(excel_path, index=False)
            print(f"Excel template created: {excel_path}")
            print("Please fill in the required data in the Excel file and save it.")

            input("Press Enter once you've completed and saved the Excel file...")

            logging.info(f"Reading data from: {excel_path}")
            df = pd.read_excel(excel_path)

            for idx, row in df.iterrows():
                logging.info(f"Generating document for row {idx + 1}...")
                self.responses = row.to_dict()
                output_path = self.output_path.replace(".docx", f"_{idx + 1}.docx")
                self.fill_document(output_path)

            print("All documents have been generated successfully.")
        except Exception as e:
            logging.error(f"Error handling multiple documents: {e}")
            raise

def parse_args():
    """ Parse command-line arguments for flexibility. """
    parser = argparse.ArgumentParser(description="Contract Generator Tool")
    parser.add_argument("--template", required=True, help="Path to the contract template file")
    parser.add_argument("--output", required=True, help="Path to save the generated contract")
    parser.add_argument("--type", required=True, help="Type of contract (e.g., NDA, SERVICE AGREEMENT)")
    return parser.parse_args()

def main():
    args = parse_args()

    contract_generator = ContractGenerator(
        template_path=args.template,
        output_path=args.output,
        contract_type=args.type
    )

    try:
        choice = input("Do you want to fill one document or many? (Type 'one' or 'many'): ").strip().lower()

        if choice == 'one':
            contract_generator.extract_placeholders()
            questions = contract_generator.generate_questions()
            contract_generator.collect_responses(questions)
            contract_generator.fill_document()
            pdf_path = args.output.replace(".docx", ".pdf")
            contract_generator.export_to_pdf(pdf_path)

        elif choice == 'many':
            contract_generator.handle_multiple_documents()

        else:
            print("Invalid choice. Please restart the program and choose 'one' or 'many'.")
    except Exception as e:
        logging.error(f"An error occurred: {e}")
        print("An error occurred during contract generation. Please check the logs for details.")

if __name__ == "__main__":
    main()