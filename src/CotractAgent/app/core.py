import logging
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from transformers import pipeline
from typing import List, Dict

# Configure logging
logging.basicConfig(
    filename="logs/contract_generator.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

class ContractGenerator:
    def __init__(self):
        try:
            logging.info("Initializing Hugging Face model...")
            self.question_generator = pipeline("text2text-generation", model="google/flan-t5-base")
        except Exception as e:
            logging.error(f"Error loading Hugging Face model: {e}")
            raise

    def extract_placeholders(self, template_path: str) -> List[str]:
        """Extract placeholders from a template document."""
        try:
            logging.info(f"Extracting placeholders from: {template_path}")
            document = Document(template_path)
            placeholders = []
            for para in document.paragraphs:
                if "[" in para.text and "]" in para.text:
                    field = para.text[para.text.find("[") + 1:para.text.find("]")]
                    placeholders.append(field)
            logging.info(f"Detected placeholders: {placeholders}")
            return placeholders
        except Exception as e:
            logging.error(f"Error extracting placeholders: {e}")
            raise

    def generate_questions(self, placeholders: List[str]) -> Dict[str, str]:
        """Generate contextual and professional questions for placeholders."""
        try:
            logging.info("Generating questions for placeholders...")
            questions = {}
            for placeholder in placeholders:
                context_question = self._generate_contextual_question(placeholder)

                if context_question == "FALLBACK":
                    input_text = f"Generate a formal and professional question about the field '{placeholder}' in a contract."
                    generated_question = self.question_generator(input_text)[0]['generated_text']
                    questions[placeholder] = self._postprocess_question(generated_question)
                else:
                    questions[placeholder] = context_question
            logging.info("Questions generated successfully.")
            return questions
        except Exception as e:
            logging.error(f"Error generating questions: {e}")
            raise

    def _generate_contextual_question(self, placeholder: str) -> str:
        """Return predefined questions or fallback to dynamic generation."""
        placeholder = placeholder.strip().upper()
        question_map = {
            'DATE': "What is the effective date of the contract?",
            'DISCLOSING_PARTY_NAME': "Who is the disclosing party mentioned in the contract?",
            'RECEIVING_PARTY_NAME': "Who is the receiving party of the confidential information?",
            'CONFIDENTIAL_INFO_DESCRIPTION': "What information is classified as confidential under this agreement?",
            'DURATION': "What is the duration of the agreement?"
        }
        return question_map.get(placeholder, "FALLBACK")

    def _postprocess_question(self, question: str) -> str:
        """Post-process generated questions for clarity and professionalism."""
        question = question.strip()
        if question.startswith("What is the best way to describe"):
            question = question.replace("What is the best way to describe", "What is")
        question = question[0].capitalize() + question[1:]  # Capitalize first letter
        return question

    def fill_document(self, template_path: str, output_path: str, responses: Dict[str, str]):
        """Fill placeholders in a document with user responses."""
        try:
            logging.info(f"Filling placeholders in: {template_path}")
            document = Document(template_path)
            for para in document.paragraphs:
                for field, response in responses.items():
                    if f"[{field}]" in para.text:
                        para.text = para.text.replace(f"[{field}]", str(response))
            document.save(output_path)
            logging.info(f"Document saved at: {output_path}")
        except Exception as e:
            logging.error(f"Error filling document: {e}")
            raise

    def export_to_pdf(self, pdf_path: str, responses: Dict[str, str]):
        """Export responses to a PDF file."""
        try:
            logging.info(f"Exporting responses to PDF: {pdf_path}")
            c = canvas.Canvas(pdf_path, pagesize=letter)
            c.setFont("Helvetica-Bold", 14)
            c.drawString(100, 750, "User Responses for Generated Contract")
            c.setFont("Helvetica", 12)

            y_pos = 730
            for field, response in responses.items():
                if y_pos < 50:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y_pos = 750
                c.drawString(100, y_pos, f"{field}: {response}")
                y_pos -= 20
            c.save()
            logging.info(f"PDF exported successfully to {pdf_path}")
        except Exception as e:
            logging.error(f"Error exporting to PDF: {e}")
            raise