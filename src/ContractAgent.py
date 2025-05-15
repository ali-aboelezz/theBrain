import os
import logging
import tkinter as tk
from tkinter import filedialog
import pandas as pd
from docx import Document
from transformers import pipeline
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Configure logging
logging.basicConfig(
    filename="contract_generator.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

class ContractGenerator:
    def __init__(self, template_path, output_path, contract_type):
        self.template_path = template_path
        self.output_path = output_path
        self.contract_type = contract_type
        self.placeholders = []
        self.responses = {}

        try:
            logging.info("Loading Hugging Face Flan-T5 question generator model...")
            self.question_generator = pipeline("text2text-generation", model="google/flan-t5-base")
        except Exception as e:
            logging.error(f"Error loading model: {e}")
            raise

    def extract_placeholders(self):
        try:
            logging.info(f"Extracting placeholders from: {self.template_path}")
            document = Document(self.template_path)
            for para in document.paragraphs:
                if "[" in para.text and "]" in para.text:
                    field = para.text[para.text.find("[") + 1:para.text.find("]")]
                    self.placeholders.append(field)
            logging.info(f"Detected placeholders: {self.placeholders}")
        except Exception as e:
            logging.error(f"Error extracting placeholders: {e}")
            raise

    def generate_questions(self):
        try:
            logging.info("Generating questions for placeholders...")
            questions = {}
            for placeholder in self.placeholders:
                context_question = self._generate_contextual_question(placeholder)
                if context_question == "FALLBACK":
                    input_text = f"Generate a formal and professional question in English about the field '{placeholder}' in a contract."
                    generated_question = self.question_generator(input_text)[0]['generated_text']
                    questions[placeholder] = self._postprocess_question(generated_question)
                else:
                    questions[placeholder] = context_question
            logging.info("Generated questions successfully.")
            return questions
        except Exception as e:
            logging.error(f"Error generating questions: {e}")
            raise

    def _generate_contextual_question(self, placeholder):
        placeholder = placeholder.strip().upper()
        question_map = {
            'DATE': "What is the effective date of the contract?",
            'DISCLOSING_PARTY_NAME': "Who is the disclosing party mentioned in the contract?",
            'RECEIVING_PARTY_NAME': "Who is the receiving party of the confidential information?",
            'CONFIDENTIAL_INFO_DESCRIPTION': "What information is classified as confidential under this agreement?",
            'DURATION': "What is the duration of the agreement?"
        }
        return question_map.get(placeholder, "FALLBACK")

    def _postprocess_question(self, question):
        question = question.strip()
        if question.startswith("What is the best way to describe"):
            question = question.replace("What is the best way to describe", "What is")
        return question[0].capitalize() + question[1:] if question else question

    def collect_responses(self, questions):
        try:
            logging.info("Collecting responses from the user...")
            for field, question in questions.items():
                while True:
                    user_input = input(f"{question}: ").strip()
                    if user_input:
                        self.responses[field] = user_input
                        break
                    else:
                        print("Input cannot be empty. Please try again.")
            logging.info(f"Collected responses: {self.responses}")
        except Exception as e:
            logging.error(f"Error collecting user responses: {e}")
            raise

    def fill_document(self, custom_output_path=None):
        try:
            logging.info(f"Filling placeholders in: {self.template_path}")
            document = Document(self.template_path)

            for para in document.paragraphs:
                for field, response in self.responses.items():
                    if f"[{field}]" in para.text:
                        para.text = para.text.replace(f"[{field}]", str(response))

            save_path = custom_output_path if custom_output_path else self.output_path
            document.save(save_path)
            logging.info(f"Contract saved as: {save_path}")
        except Exception as e:
            logging.error(f"Error filling document: {e}")
            raise

    def export_to_pdf(self, pdf_path):
        try:
            logging.info(f"Exporting user responses to PDF: {pdf_path}")
            c = canvas.Canvas(pdf_path, pagesize=letter)

            c.setFont("Helvetica-Bold", 14)
            c.drawString(100, 750, "User Responses for Generated Contract")
            c.setFont("Helvetica", 12)

            y_pos = 730
            for field, response in self.responses.items():
                if y_pos < 50:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y_pos = 750
                c.drawString(100, y_pos, f"{field}: {response}")
                y_pos -= 20

            c.save()
            logging.info(f"PDF exported successfully to {pdf_path}")
        except Exception as e:
            logging.error(f"Error exporting to PDF: {e}")
            raise

    def handle_multiple_documents(self):
        try:
            logging.info("Handling multiple document generation...")
            self.extract_placeholders()

            excel_path = "contract_data.xlsx"
            logging.info(f"Creating Excel template at: {excel_path}")
            df = pd.DataFrame(columns=self.placeholders)
            df.to_excel(excel_path, index=False)
            print(f"📄 Excel template created: {excel_path}")
            print("✍️ Please fill in the required data in the Excel file and save it.")

            input("🔔 Press Enter once you've completed and saved the Excel file...")

            logging.info(f"Reading data from: {excel_path}")
            df = pd.read_excel(excel_path)

            for idx, row in df.iterrows():
                logging.info(f"Generating document for row {idx + 1}...")
                self.responses = row.to_dict()
                output_path = self.output_path.replace(".docx", f"_{idx + 1}.docx")
                self.fill_document(output_path)

            print("✅ All documents have been generated successfully.")
        except Exception as e:
            logging.error(f"Error handling multiple documents: {e}")
            raise

def pick_template_file():
    """GUI to select a DOCX template, with forced window pop-up."""
    root = tk.Tk()
    root.lift()
    root.attributes('-topmost', True)
    root.withdraw()
    root.update()

    file_path = filedialog.askopenfilename(
        title="Select a Word Template",
        filetypes=[("Word Documents", "*.docx")]
    )

    root.destroy()
    return file_path

def main():
    print("👋 Welcome to the Contract Generator Assistant!")

    contract_type = input("📝 What type of document would you like to create (e.g., NDA, Service Agreement)? ").strip()

    print("📂 Please select your DOCX template file...")
    template_path = pick_template_file()

    if not template_path or not template_path.lower().endswith(".docx"):
        print("❗ No valid file selected. Exiting.")
        return

    output_filename = input("💾 Enter the name for the generated DOCX file (e.g., contract.docx): ").strip()
    if not output_filename.lower().endswith(".docx"):
        output_filename += ".docx"

    output_path = os.path.join(os.getcwd(), output_filename)

    contract_generator = ContractGenerator(
        template_path=template_path,
        output_path=output_path,
        contract_type=contract_type
    )

    try:
        choice = input("🔄 Do you want to create one document or many? (Type 'one' or 'many'): ").strip().lower()

        if choice == "one":
            contract_generator.extract_placeholders()
            questions = contract_generator.generate_questions()
            contract_generator.collect_responses(questions)
            contract_generator.fill_document()
            pdf_path = output_path.replace(".docx", ".pdf")
            contract_generator.export_to_pdf(pdf_path)
            print("✅ Single document created successfully.")

        elif choice == "many":
            contract_generator.handle_multiple_documents()

        else:
            print("❗ Invalid choice. Please restart and type 'one' or 'many'.")

    except Exception as e:
        logging.error(f"An error occurred: {e}")
        print("❌ An error occurred during contract generation. Please check the logs for details.")

if __name__ == "__main__":
    main()
